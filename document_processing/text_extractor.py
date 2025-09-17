"""
Text Extractor Module

Handles PDF and DOCX text extraction using Docling's advanced capabilities.
"""

import re
from pathlib import Path
from typing import List, Dict, Any, Tuple

from docling.document_converter import DocumentConverter
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.datamodel.base_models import InputFormat


class DocumentTextExtractor:
    """Extract text from PDF and DOCX documents using Docling"""
    
    def __init__(self):
        self.docx_processor = None
        self.converter = None
        self.pipeline_options = None
        self.pipeline_info = None
        self._setup_docling_converter()
        self._setup_docx_processor()
    
    def _setup_docx_processor(self):
        """Setup enhanced DOCX processing with python-docx"""
        try:
            import docx
            self.docx_processor = docx
            print("[i] Enhanced DOCX processing available via python-docx")
        except ImportError:
            print("[i] Enhanced DOCX processing unavailable - using standard processing")
    
    def _setup_docling_converter(self):
        """Configure Docling for maximum form scanning accuracy"""
        # Configure pipeline for maximum accuracy
        self.pipeline_options = PdfPipelineOptions()
        self.pipeline_options.do_ocr = True  # Enable OCR for scanned forms
        self.pipeline_options.do_table_structure = True  # Detect table structures
        self.pipeline_options.images_scale = 2.0  # Higher resolution for better OCR
        self.pipeline_options.generate_page_images = False  # Don't need page images
        self.pipeline_options.generate_table_images = False  # Don't need table images
        self.pipeline_options.generate_picture_images = False  # Don't need picture images
        
        # Force full page OCR for maximum field detection
        if hasattr(self.pipeline_options.ocr_options, 'force_full_page_ocr'):
            self.pipeline_options.ocr_options.force_full_page_ocr = True
        
        # Create converter with optimized settings
        self.converter = DocumentConverter()
        
        # Store pipeline info for reporting
        self.pipeline_info = {
            'pipeline': 'StandardPdfPipeline',
            'backend': 'DoclingParseDocumentBackend', 
            'ocr_enabled': self.pipeline_options.do_ocr,
            'ocr_engine': 'EasyOCR',  # Docling's default OCR engine
            'table_structure': self.pipeline_options.do_table_structure,
            'images_scale': self.pipeline_options.images_scale
        }
        
    def remove_practice_headers_footers(self, text_lines: List[str]) -> List[str]:
        """Universal header/footer removal to clean practice information from consent forms"""
        
        # Define patterns for practice information that should be removed
        practice_patterns = [
            # Email and website patterns
            r'.*@.*\.com.*',
            r'.*www\..*\.com.*',
            
            # Phone number patterns
            r'.*\(\d{3}\).*\d{3}.*\d{4}.*',
            r'.*\d{3}[-\.]\d{3}[-\.]\d{4}.*',
            
            # Address patterns (street numbers, zip codes)
            r'.*\d+\s+[A-Za-z\s]+(St|Street|Ave|Avenue|Rd|Road|Dr|Drive|Blvd|Boulevard).*',
            r'.*\b\d{5}(-\d{4})?\b.*',  # ZIP codes
            
            # Practice name patterns (common dental practice words)
            r'.*(dental|dentistry|oral|surgery|care|clinic|center|associates|group|practice).*',
            
            # Footer information
            r'.*page\s+\d+.*',
            r'.*©.*\d{4}.*',
            r'.*all\s+rights\s+reserved.*',
            
            # Form metadata
            r'.*form\s*(id|number|version).*',
            r'.*revised.*\d{4}.*',
        ]
        
        # Compile patterns for efficiency
        compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in practice_patterns]
        
        # Filter out lines matching practice patterns
        filtered_lines = []
        for line in text_lines:
            line = line.strip()
            if not line:  # Keep empty lines for structure
                filtered_lines.append(line)
                continue
                
            # Check if line matches any practice pattern
            is_practice_info = any(pattern.search(line) for pattern in compiled_patterns)
            
            # Additional checks for specific practice content
            if not is_practice_info:
                line_lower = line.lower()
                # Remove lines that are clearly practice-specific
                practice_keywords = [
                    'smile solutions', 'dental office', 'family dentistry', 
                    'cosmetic dentistry', 'orthodontics', 'endodontics',
                    'periodontics', 'oral surgery', 'implant dentistry'
                ]
                is_practice_info = any(keyword in line_lower for keyword in practice_keywords)
            
            if not is_practice_info:
                filtered_lines.append(line)
        
        return filtered_lines

    def extract_enhanced_docx_structure(self, document_path: Path) -> Tuple[List[str], Dict[str, Any]]:
        """Enhanced DOCX structure recognition using python-docx"""
        if not self.docx_processor:
            # Fallback to standard processing
            return self.extract_text_from_document(document_path)
        
        try:
            doc = self.docx_processor.Document(document_path)
            text_lines = []
            
            # Extract paragraphs with enhanced structure detection
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Detect if this is a form field line
                    if any(char in text for char in ['_', '□', '☐', '◻', '■', '☑', '✓']):
                        # This might be a form field
                        text_lines.append(text)
                    elif text.endswith(':'):
                        # This might be a field label
                        text_lines.append(text)
                    else:
                        # Regular content
                        text_lines.append(text)
                else:
                    # Preserve structure with empty lines
                    text_lines.append("")
            
            # Extract tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_lines.append(" | ".join(row_text))
            
            pipeline_info = {
                "pipeline": "SimplePipeline", 
                "backend": "python-docx",
                "format": "DOCX"
            }
            
            print(f"[i] Enhanced DOCX extraction: {len(text_lines)} lines from {document_path.name}")
            return text_lines, pipeline_info
            
        except Exception as e:
            print(f"[!] Enhanced DOCX processing failed: {e}, falling back to standard processing")
            return self.extract_text_from_document(document_path)

    def extract_text_from_document(self, document_path: Path) -> Tuple[List[str], Dict[str, Any]]:
        """Extract text from PDF or DOCX using enhanced capabilities"""
        document_path = Path(document_path)
        
        if not document_path.exists():
            raise ValueError(f"Document not found: {document_path}")
        
        # For DOCX files, try enhanced processing first
        if document_path.suffix.lower() in ['.docx', '.doc'] and self.docx_processor:
            return self.extract_enhanced_docx_structure(document_path)
        
        try:
            # Convert document using Docling (supports PDF, DOCX, and other formats)
            result = self.converter.convert(str(document_path))
            
            # Extract text with superior layout preservation - use markdown for checkbox preservation
            full_text = result.document.export_to_markdown()  # Changed from export_to_text()
            all_lines = full_text.split('\n')
            # Keep empty lines for proper question-option proximity in radio detection, but limit empty line runs
            text_lines = []
            for line in all_lines:
                stripped = line.strip()
                # Keep line but avoid excessive empty line runs
                if stripped or (text_lines and text_lines[-1].strip()):
                    text_lines.append(stripped)
            
            # UNIVERSAL HEADER/FOOTER REMOVAL for all document types
            text_lines = self.remove_practice_headers_footers(text_lines)
            
            # Update pipeline info with actual conversion details
            pipeline_info = self.pipeline_info.copy()
            pipeline_info['document_name'] = result.document.name
            pipeline_info['elements_extracted'] = len(list(result.document.texts))
            
            # Detect document format
            document_suffix = document_path.suffix.lower()
            if document_suffix in ['.docx', '.doc']:
                pipeline_info['document_format'] = 'DOCX'
                pipeline_info['ocr_used'] = False  # DOCX doesn't need OCR
            else:
                pipeline_info['document_format'] = 'PDF'
                pipeline_info['ocr_used'] = pipeline_info.get('ocr_enabled', False)
            
            return text_lines, pipeline_info
            
        except Exception as e:
            print(f"Error reading document {document_path} with Docling: {e}")
            return [], self.pipeline_info